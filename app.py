from flask import Flask, render_template, request, send_file, after_this_request,redirect
from flask_bootstrap import Bootstrap5
from docx import Document
from python_docx_replace import docx_replace
import os
from datetime import datetime


app = Flask(__name__)
app.config['MODULI_FOLDER'] = 'module'
app.config['OUTPUT_FOLDER'] = 'output'
bootstrap = Bootstrap5(app)

# Assicurati che la cartella di output esista
if not os.path.exists(app.config['OUTPUT_FOLDER']):
    os.makedirs(app.config['OUTPUT_FOLDER'])

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/fviaggio', methods=['GET','POST'])
def fviaggio():
    form = request.form
    
    if request.method == 'POST':
        # Ottieni i dati dal modulo
        mydict = {
            'qual': request.form['qual'],
            'name': request.form['name'],
            'surname': request.form['surname'],
            'perid': request.form['perid'],
            'nr_fvg': request.form['nr_fvg'],
            'motive': request.form['motive'],
            'cod': request.form['mission_code'],
            'cod_off': request.form['office_code'],
            'dest': request.form['destination'],
            'data': datetime.strptime(request.form['date'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'month': datetime.strptime(request.form['date'],'%Y-%m-%d').strftime('%m'),
            'year': datetime.strptime(request.form['date'],'%Y-%m-%d').strftime('%Y'),
            'h_serv': request.form['h_serv'],
            'data_fvg': datetime.strptime(request.form['data_fvg'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'from1': request.form['from1'],
            'to1': request.form['to1'],
            'from2': request.form['from2'],
            'to2': request.form['to2'],
            'from3': request.form['from3'],
            'to3': request.form['to3'],
            'from4': request.form['from4'],
            'to4': request.form['to4'],
            'from5': request.form['from5'],
            'to5': request.form['to5'],
            'from6': request.form['from6'],
            'to6': request.form['to6'],
            'yn1': '[X]' if request.form['canteen'] == 'YES' else '[ ]',
            'yn2': '[X]' if request.form['canteen'] == 'NO' else '[ ]',
            'yn3': '[X]' if request.form['lodge'] == 'YES' else '[ ]',
            'yn4': '[X]' if request.form['lodge'] == 'NO' else '[ ]',
            'canteen': 'PRESENTE' if request.form['canteen'] == 'YES' else 'NON PRESENTE',
            'accomodation': 'PRESENTE' if request.form['lodge'] == 'YES' else 'NON PRESENTE'
            }

        # Carica il template Word
        template_path = os.path.join(app.config['MODULI_FOLDER'], 'f_viaggio.docx')
        doc = Document(template_path)

        # Sostituisci i segnaposto con i dati del modulo
        docx_replace(doc, **mydict)
        # Salva il documento modificato 
        output_filename = "{}_{}_foglio_viaggio_{}.docx".format(mydict['surname'], mydict['name'], datetime.now().strftime("%Y%m%d_%H%M%S"))
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        doc.save(output_path)

        return redirect('/download/{}'.format(output_filename))
    
    # Se non è una richiesta POST, mostra il modulo 
    return render_template('fviaggio.html', form=form)

@app.route('/straordinario', methods=['GET','POST'])
def straordinario():
    form = request.form
    filepaths = []
    
    if request.method == 'POST':
        # Ottieni i dati dal modulo
        mydict = {
            'qual': request.form['qual'],
            'surname': request.form['surname'],
            'name': request.form['name'],
            'subject': "{} {}".format(request.form['surname'],request.form['name']),
            'perid': request.form['perid'],
            'datastr': datetime.strptime(request.form['datastr'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'working_hours': request.form['working_hours'],
            'h_for': request.form['h_for'],
            'h_to': request.form['h_to'],
            'motive': request.form['motive'],
            'h_str' : request.form['h_str'],
            'h_str_fn': request.form['h_str_fn'],
            'h_str_nf': request.form['h_str_nf'],
            'data' : datetime.now().strftime('%d-%m-%Y'),
            'cash': '[x]' if (request.form['cash'] == 'cash') else '[ ]',
            'comp': '[ ]' if (request.form['cash'] == 'cash') else '[x]',
            'datanow': datetime.now().strftime('%d-%m-%Y')
           }
        
        # crea modulo autorizzazione
        if 'mod_aut' in request.form:
            # Carica il template Word
            template_path = os.path.join(app.config['MODULI_FOLDER'], 'str_autorizzazione.docx')
            doc = Document(template_path)

            # Sostituisci i segnaposto con i dati del modulo
            docx_replace(doc, **mydict)
            # Salva il documento modificato 
            output_filename = "{}_{}_{}_autorizzazione_straordinario_{}.docx".format(mydict['surname'], mydict['name'], mydict['perid'],mydict['datastr'])
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            doc.save(output_path)

            filepaths.append({'filename': output_filename, 'display_name': 'Autrorizzazione straordinario'})
    
        # crea modulo dichiarazione
        if  'mod_dic' in request.form:
            # Carica il template Word
            template_path = os.path.join(app.config['MODULI_FOLDER'], 'str_dichiarazione.docx')
            doc = Document(template_path)

            # Sostituisci i segnaposto con i dati del modulo
            docx_replace(doc, **mydict)
            # Salva il documento modificato 
            output_filename = "{}_{}_{}_dichiarazione_straordinario_{}.docx".format(mydict['surname'], mydict['name'], mydict['perid'],mydict['datastr'])
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            doc.save(output_path)

            filepaths.append({'filename': output_filename, 'display_name': 'Dichiarazione straordinario'})

        # crea modulo ratifica
        if 'mod_rat' in request.form:
            # Carica il template Word
            template_path = os.path.join(app.config['MODULI_FOLDER'], 'str_ratifica.docx')
            doc = Document(template_path)

            # Sostituisci i segnaposto con i dati del modulo
            docx_replace(doc, **mydict)
            # Salva il documento modificato 
            output_filename = "{}_{}_{}_ratifica_straordinario_{}.docx".format(mydict['surname'], mydict['name'], mydict['perid'],mydict['datastr'])
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            doc.save(output_path)

            filepaths.append({'filename': output_filename, 'display_name': 'Ratifica straordinario'})

        return render_template('straordinario.html',files=filepaths)
    
    # Se non è una richiesta POST, mostra il modulo 
    return render_template('straordinario.html', form=form, files=filepaths)

@app.route('/elearning', methods=['GET','POST'])
def elearning():
    form = request.form
    filepaths = []

    # Se la richiesta è di tipo POST, elabora i dati del modulo
    # e genera il documento Word
    
    if request.method == 'POST':
        # Ottieni i dati dal modulo
        mydict = {
            'year': datetime.strptime(request.form['date_agg'],'%Y-%m-%d').strftime('%Y'),
            'qual': request.form['qual'],
            'surname': request.form['surname'],
            'name': request.form['name'],
            'perid': request.form['perid'],
            'module': request.form['module'],
            'nr_day': request.form['nr_day'],
            'date_agg': datetime.strptime(request.form['date_agg'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'h_agg': request.form['h_agg'],
            'data_old_agg': datetime.strptime(request.form['data_old_agg'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'home': '[X]' if request.form['home'] == 'home' else '[ ]',
            'office': '[X]' if request.form['home'] == 'office' else '[ ]',
            'address': request.form['address'],
            'date_now': datetime.now().strftime('%d-%m-%Y'),
            'type': request.form['type'],
            'where': 'il proprio domicilio' if request.form['home'] == 'home' else 'il proprio ufficio'
        }

        # Autorizzazione e-learning
        if 'auth' in request.form:
            
            # Carica il template Word
            template_path = os.path.join(app.config['MODULI_FOLDER'], 'richiesta_e_learning.docx')
            doc = Document(template_path)

            # Sostituisci i segnaposto con i dati del modulo
            docx_replace(doc, **mydict)
            # Salva il documento modificato 
            output_filename = "{}_{}_{}_autorizzazione_elearning_{}.docx".format(mydict['surname'], mydict['name'],mydict['perid'], mydict['date_agg'])
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            doc.save(output_path)

            filepaths.append({'filename': output_filename, 'display_name': 'Autorizzazione e-learning'})
        
        # Dichiarazione e-learning
        if 'decl' in request.form:
            # Carica il template Word
            template_path = os.path.join(app.config['MODULI_FOLDER'], 'autocertificazione_e_learning.docx')
            doc = Document(template_path)

            # Sostituisci i segnaposto con i dati del modulo
            docx_replace(doc, **mydict)
            # Salva il documento modificato 
            output_filename = "{}_{}_{}_autocertificazione_elearning_{}.docx".format(mydict['surname'], mydict['name'],mydict['perid'], mydict['date_agg'])
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            doc.save(output_path)

            filepaths.append({'filename': output_filename, 'display_name': 'autocertificazione e-learning'})

    
    # Se non è una richiesta POST, mostra il modulo 
    return render_template('elearning.html', form=form, files=filepaths)

@app.route('/cstrmalattia',methods=['GET','POST'])
def cstrmalattia():
    form = request.form

    if request.method == 'POST':
        # Ottieni i dati dal modulo
        mydict = {
            'qual': request.form['qual'],
            'surname': request.form['surname'],
            'name': request.form['name'],
            'perid': request.form['perid'],
            'old_cstr': '[X]' if 'old_cstr' in request.form else '[ ]',
            'date_old_cstr': datetime.strptime(request.form['date_old_cstr'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'old_asp': '[X]' if 'old_asp' in request.form else '[ ]',
            'date_old_asp': datetime.strptime(request.form['date_old_asp'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'cstr': '[X]' if 'cstr' in request.form  else '[ ]',
            'cstr_from': datetime.strptime(request.form['cstr_from'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'cstr_to': datetime.strptime(request.form['cstr_to'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'cstr_days': request.form['cstr_days'],
            'asp': '[x]' if 'asp' in request.form else '[ ]',
            'asp_from': datetime.strptime(request.form['asp_from'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'asp_to': datetime.strptime(request.form['asp_to'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'asp_days': request.form['asp_days'],
            'address': request.form['address'],
            'tel': request.form['tel'],
            'date_now': datetime.now().strftime('%d-%m-%Y')
           }
                
        # Carica il template Word
        template_path = os.path.join(app.config['MODULI_FOLDER'], 'cs_malattia.docx')
        doc = Document(template_path)

        # Sostituisci i segnaposto con i dati del modulo
        docx_replace(doc, **mydict)
        # Salva il documento modificato 
        output_filename = "{}_{}_{}_cstr_malattia_aspettativa_{}.docx".format(mydict['surname'], mydict['name'], mydict['perid'],mydict['date_now'])
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        doc.save(output_path)
        
        return redirect('/download/{}'.format(output_filename))
    
    # Se non è una richiesta POST, mostra il modulo
    return render_template('cstrmalattia.html', form=form)

@app.route('/dct',methods=['GET','POST'])
def dct():
    form = request.form

    if request.method == 'POST':
        # Ottieni i dati dal modulo
        mydict = {
            'qual': request.form['qual'],
            'surname': request.form['surname'],
            'name': request.form['name'],
            'perid': request.form['perid'],
            'dec': request.form['dec'],
            'date': datetime.strptime(request.form['date'],'%Y-%m-%d').strftime('%d-%m-%Y'),
            'reason': request.form['reason'],
            'shift_to': request.form['shift_to'],
            'shift_from': request.form['shift_from'],
            'data_now': datetime.now().strftime('%d-%m-%Y')
           }
                
        # Carica il template Word
        template_path = os.path.join(app.config['MODULI_FOLDER'], 'cambio_turno.docx')
        doc = Document(template_path)

        # Sostituisci i segnaposto con i dati del modulo
        docx_replace(doc, **mydict)
        # Salva il documento modificato 
        output_filename = "{}_{}_{}_decreto_cambio_turno_{}.docx".format(mydict['surname'], mydict['name'], mydict['perid'],mydict['date'])
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        doc.save(output_path)
        
        return redirect('/download/{}'.format(output_filename))
    
    # Se non è una richiesta POST, mostra il modulo
    return render_template('dct.html', form=form)

@app.route('/download/<filename>')
def download(filename):
    # Percorso completo del file da scaricare
    file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)

    response = send_file(file_path, as_attachment=True)

    @after_this_request
    def delete_file(response):
        try:
            os.remove(file_path)  # Rimuovi il file dopo il download
        except Exception as e:
            print("Error removing or closing downloaded file handle", e)
        return response

    # Invia il file come download
    return response 

if __name__ == '__main__':
    app.run()
